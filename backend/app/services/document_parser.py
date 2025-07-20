"""
Document Parsing Service
Handles file upload, validation, and text extraction for policy documents
"""

import os
import magic
import fitz  # PyMuPDF
import docx
import pytesseract
from PIL import Image
import io
import re
import hashlib
from typing import Dict, List, Tuple, Optional, Any
from datetime import datetime
import logging
from pathlib import Path

# Configure logging
logger = logging.getLogger(__name__)

class DocumentParserError(Exception):
    """Custom exception for document parsing errors"""
    pass

class DocumentValidationError(DocumentParserError):
    """Exception for document validation errors"""
    pass

class DocumentParser:
    """
    Document parser service for handling PDF, DOCX, and text files
    """
    
    # Supported file types and their MIME types
    SUPPORTED_TYPES = {
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'text/plain': 'txt',
        'text/markdown': 'md'
    }
    
    # File size limits (10MB as per PRD)
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    # Maximum text length for processing
    MAX_TEXT_LENGTH = 1000000  # 1M characters
    
    def __init__(self):
        """Initialize the document parser"""
        self.upload_dir = Path("uploads")
        self.upload_dir.mkdir(exist_ok=True)
        
    def validate_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Validate uploaded file for type, size, and content
        
        Args:
            file_content: Raw file content
            filename: Original filename
            
        Returns:
            Dict containing validation results and file metadata
        """
        try:
            # Check file size
            file_size = len(file_content)
            if file_size > self.MAX_FILE_SIZE:
                raise DocumentValidationError(
                    f"File size {file_size} bytes exceeds maximum allowed size of {self.MAX_FILE_SIZE} bytes"
                )
            
            # Detect MIME type using magic bytes
            mime_type = magic.from_buffer(file_content, mime=True)
            
            if mime_type not in self.SUPPORTED_TYPES:
                raise DocumentValidationError(
                    f"Unsupported file type: {mime_type}. Supported types: {list(self.SUPPORTED_TYPES.keys())}"
                )
            
            # Generate file hash for deduplication
            file_hash = hashlib.sha256(file_content).hexdigest()
            
            # Validate filename
            if not filename or len(filename) > 255:
                raise DocumentValidationError("Invalid filename")
            
            # Check for potentially malicious extensions
            dangerous_extensions = ['.exe', '.bat', '.cmd', '.com', '.scr', '.vbs', '.js']
            file_ext = Path(filename).suffix.lower()
            if file_ext in dangerous_extensions:
                raise DocumentValidationError(f"Potentially dangerous file extension: {file_ext}")
            
            return {
                "valid": True,
                "mime_type": mime_type,
                "file_type": self.SUPPORTED_TYPES[mime_type],
                "file_size": file_size,
                "file_hash": file_hash,
                "filename": filename,
                "uploaded_at": datetime.utcnow().isoformat()
            }
            
        except DocumentValidationError:
            raise
        except Exception as e:
            logger.error(f"File validation error: {e}")
            raise DocumentValidationError(f"File validation failed: {str(e)}")
    
    def save_file(self, file_content: bytes, filename: str, file_hash: str) -> str:
        """
        Save uploaded file to disk
        
        Args:
            file_content: Raw file content
            filename: Original filename
            file_hash: File hash for deduplication
            
        Returns:
            Path to saved file
        """
        try:
            # Create filename with hash to avoid conflicts
            file_ext = Path(filename).suffix
            safe_filename = f"{file_hash}{file_ext}"
            file_path = self.upload_dir / safe_filename
            
            # Save file
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            logger.info(f"File saved: {file_path}")
            return str(file_path)
            
        except Exception as e:
            logger.error(f"File save error: {e}")
            raise DocumentParserError(f"Failed to save file: {str(e)}")
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file using PyMuPDF
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            doc = fitz.Document(file_path)
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Try text extraction first
                page_text = page.get_text()
                
                # If no text found, try OCR
                if not page_text.strip():
                    logger.info(f"No text found on page {page_num + 1}, attempting OCR")
                    pix = page.get_pixmap()
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    page_text = pytesseract.image_to_string(img)
                
                text += page_text + "\n"
            
            doc.close()
            
            # Clean up extracted text
            text = self._clean_text(text)
            
            if len(text) > self.MAX_TEXT_LENGTH:
                logger.warning(f"Text length {len(text)} exceeds maximum, truncating")
                text = text[:self.MAX_TEXT_LENGTH]
            
            return text
            
        except Exception as e:
            logger.error(f"PDF text extraction error: {e}")
            raise DocumentParserError(f"Failed to extract text from PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file using python-docx
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            
            # Clean up extracted text
            text = self._clean_text(text)
            
            if len(text) > self.MAX_TEXT_LENGTH:
                logger.warning(f"Text length {len(text)} exceeds maximum, truncating")
                text = text[:self.MAX_TEXT_LENGTH]
            
            return text
            
        except Exception as e:
            logger.error(f"DOCX text extraction error: {e}")
            raise DocumentParserError(f"Failed to extract text from DOCX: {str(e)}")
    
    def extract_text_from_txt(self, file_content: bytes) -> str:
        """
        Extract text from plain text file
        
        Args:
            file_content: Raw file content
            
        Returns:
            Extracted text
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                # If all encodings fail, use utf-8 with error handling
                text = file_content.decode('utf-8', errors='replace')
            
            # Clean up extracted text
            text = self._clean_text(text)
            
            if len(text) > self.MAX_TEXT_LENGTH:
                logger.warning(f"Text length {len(text)} exceeds maximum, truncating")
                text = text[:self.MAX_TEXT_LENGTH]
            
            return text
            
        except Exception as e:
            logger.error(f"Text file extraction error: {e}")
            raise DocumentParserError(f"Failed to extract text from file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove common header/footer patterns
        text = re.sub(r'Page \d+ of \d+', '', text)
        text = re.sub(r'Confidential|Internal Use Only|Draft', '', text)
        
        # Remove email headers and footers
        text = re.sub(r'From:.*?\n', '', text)
        text = re.sub(r'Sent:.*?\n', '', text)
        text = re.sub(r'To:.*?\n', '', text)
        text = re.sub(r'Subject:.*?\n', '', text)
        
        # Remove common document artifacts
        text = re.sub(r'\[.*?\]', '', text)  # Remove bracketed content
        text = re.sub(r'^\s*[-_=]+\s*$', '', text, flags=re.MULTILINE)  # Remove separator lines
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract text from file based on type
        
        Args:
            file_path: Path to file
            file_type: Type of file (pdf, docx, txt)
            
        Returns:
            Extracted text
        """
        if file_type == 'pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            return self.extract_text_from_docx(file_path)
        elif file_type in ['txt', 'md']:
            with open(file_path, 'rb') as f:
                return self.extract_text_from_txt(f.read())
        else:
            raise DocumentParserError(f"Unsupported file type: {file_type}")
    
    def process_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Complete document processing pipeline
        
        Args:
            file_content: Raw file content
            filename: Original filename
            
        Returns:
            Processing results with metadata and extracted text
        """
        start_time = datetime.utcnow()
        try:
            # Step 1: Validate file
            validation_result = self.validate_file(file_content, filename)
            
            # Step 2: Save file
            file_path = self.save_file(file_content, filename, validation_result['file_hash'])
            
            # Step 3: Extract text
            extracted_text = self.extract_text(file_path, validation_result['file_type'])
            
            # Step 4: Generate processing metadata
            end_time = datetime.utcnow()
            processing_metadata = {
                "text_length": len(extracted_text),
                "word_count": len(extracted_text.split()),
                "line_count": len(extracted_text.split('\n')),
                "character_count": len(extracted_text),
                "processing_time_ms": int((end_time - start_time).total_seconds() * 1000),
                "file_path": file_path
            }
            
            return {
                "success": True,
                "validation": validation_result,
                "extracted_text": extracted_text,
                "processing_metadata": processing_metadata
            }
            
        except Exception as e:
            logger.error(f"Document processing error: {e}")
            return {
                "success": False,
                "error": str(e),
                "validation": validation_result if 'validation_result' in locals() else None
            }
    
    def cleanup_file(self, file_path: str) -> bool:
        """
        Clean up temporary file
        
        Args:
            file_path: Path to file to delete
            
        Returns:
            True if successful
        """
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"Cleaned up file: {file_path}")
                return True
            return False
        except Exception as e:
            logger.error(f"File cleanup error: {e}")
            return False 